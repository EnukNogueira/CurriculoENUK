from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Criação do documento
doc = Document()

# Estilo BASICO dos paragrafos
style = doc.styles['Normal']
font = style.font
font.size = Pt(12)  # DEIXA O TAMANHO DA FONTE EM 12

# CONTATO
doc.add_heading('ENUK DOS SANTOS ALVES NOGUEIRA', level=1)
doc.add_paragraph("Guarulhos, SP | (11) 97073-0434 | enuk.santos@gmail.com")
p = doc.add_paragraph()
p.add_run("LinkedIn: ").bold = True
# HYPERLINK:
p.add_run('https://www.linkedin.com/in/enuk-nogueira').hyperlink = 'https://www.linkedin.com/in/enuk-nogueira'

# ESPAÇO
doc.add_paragraph() 

# RESUMO PROFISSIONAL
doc.add_heading('RESUMO PROFISSIONAL', level=2)
doc.add_paragraph(
    "Graduando em Engenharia de Software com experiência em documentação técnica, "
    "gestão de processos e suporte técnico. Possuo conhecimento em linguagens de programação "
    "(Python, C#, Java, JavaScript), redes e protocolos de comunicação e uso de ferramentas como "
    "Jira e Power BI. Busco uma oportunidade para aplicar e expandir meus conhecimentos em TI, "
    "contribuindo para projetos inovadores e eficientes."
)

# MAIS ESPAÇO
doc.add_paragraph()

# Formação Acadêmica
doc.add_heading('FORMAÇÃO ACADÊMICA', level=2)
doc.add_paragraph("Engenharia de Software – Universidade Cruzeiro do Sul (08/2024 - 12/2026)")
doc.add_paragraph("Técnico em Administração – Universidade Eniac (02/2016 - 12/2017)")
doc.add_paragraph("Técnico em Logística – Universidade Eniac (12/2017 - 12/2018)")
doc.add_paragraph("CAI em Eletricista de Manutenção Eletroeletrônica – Senai (02/2023 - 12/2024)")

# AINDA MAIS ESPAÇO
doc.add_paragraph()

# Experiência Profissional
doc.add_heading('EXPERIÊNCIA PROFISSIONAL', level=2)
doc.add_paragraph("Phibro Animal Health – Aprendiz de Eletricista de Manutenção Eletroeletrônica (02/2023 – 12/2024)")
doc.add_paragraph("- Organização e controle de documentação técnica, ordens de serviço e laudos elétricos.")
doc.add_paragraph("- Suporte no planejamento e execução de manutenções preventivas e corretivas.")
doc.add_paragraph("- Monitoramento de cronogramas e indicadores de desempenho.")
doc.add_paragraph("- Gestão de compras e controle de materiais elétricos.")
doc.add_paragraph("- Suporte técnico e atendimento ao cliente interno.")

# DE NOVO AINDA MAIS ESPAÇO
doc.add_paragraph()

# Habilidades e Competências
doc.add_heading('HABILIDADES E COMPETÊNCIAS', level=2)
doc.add_paragraph("Linguagens de Programação: Python, C#, Java, HTML, CSS, JavaScript")
doc.add_paragraph("Desenvolvimento e Gestão: Controle de Documentação, Planejamento Estratégico, Gestão de Processos")
doc.add_paragraph("Ferramentas: Jira, Excel Avançado, Power BI, AutoCAD")
doc.add_paragraph("Infraestrutura e Redes: Redes e Protocolos de Comunicação (Básico)")

# Adicionando mais espaçamento
doc.add_paragraph()

# Certificações
doc.add_heading('CERTIFICAÇÕES', level=2)
doc.add_paragraph("Desenvolvimento Web Completo – Udemy")
doc.add_paragraph("Python Básico ao Avançado – Udemy")
doc.add_paragraph("Python Completo – Danki Code")
doc.add_paragraph("Lógica de Programação – Senai")
doc.add_paragraph("Excel Avançado – Fundação Bradesco")
doc.add_paragraph("Desenho Técnico e Sistemas Eletrônicos Digitais – AutoDesk")
doc.add_paragraph("Adobe Photoshop, Illustrator e Dreamweaver – Colégio Serrano Guardia")

# NOVAMENTE MAIS ESPAÇO
doc.add_paragraph()

# Idiomas
doc.add_heading('IDIOMAS', level=2)
doc.add_paragraph("Inglês: Intermediário (Leitura e escrita intermediária, conversação básica)")

# DE NOVO... ESPAÇO
doc.add_paragraph()

# Informações Adicionais
doc.add_heading('INFORMAÇÕES ADICIONAIS', level=2)
doc.add_paragraph("Interesse e aprendizado contínuo em TI, programação e desenvolvimento de sistemas.")

# Salvar o documento em formato .docx
doc.save('Curriculo_Enuk.docx')